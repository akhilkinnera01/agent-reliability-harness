"""
ARH Document Loader

Supports loading multiple document formats:
- PDF (.pdf)
- Word (.docx, .doc)
- EPUB (.epub)
- Markdown (.md)
- Plain text (.txt)
- HTML (.html, .htm)
- reStructuredText (.rst)
"""

import os
from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass


@dataclass
class LoadedDocument:
    """Represents a loaded document."""
    content: str
    filename: str
    format: str
    pages: int = 1
    word_count: int = 0
    char_count: int = 0


class DocumentLoader:
    """
    Universal document loader supporting multiple formats.
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    - EPUB (.epub)
    - Markdown (.md)
    - Plain text (.txt)
    - HTML (.html, .htm)
    - reStructuredText (.rst)
    - Python, JavaScript, etc. (any text-based code)
    """
    
    # Format handlers
    SUPPORTED_FORMATS = {
        ".pdf": "PDF",
        ".docx": "Word",
        ".doc": "Word (Legacy)",
        ".epub": "EPUB",
        ".md": "Markdown",
        ".txt": "Plain Text",
        ".html": "HTML",
        ".htm": "HTML",
        ".rst": "reStructuredText",
        ".py": "Python",
        ".js": "JavaScript",
        ".ts": "TypeScript",
        ".json": "JSON",
        ".xml": "XML",
        ".yaml": "YAML",
        ".yml": "YAML",
    }
    
    @classmethod
    def load(cls, file_path: str) -> LoadedDocument:
        """
        Load a document from any supported format.
        
        Args:
            file_path: Path to the document
            
        Returns:
            LoadedDocument with extracted text content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower()
        
        # Route to appropriate handler
        if ext == ".pdf":
            content, pages = cls._load_pdf(path)
        elif ext == ".docx":
            content, pages = cls._load_docx(path)
        elif ext == ".epub":
            content, pages = cls._load_epub(path)
        elif ext in [".html", ".htm"]:
            content, pages = cls._load_html(path)
        else:
            # Default: treat as plain text
            content, pages = cls._load_text(path)
        
        format_name = cls.SUPPORTED_FORMATS.get(ext, "Text")
        
        return LoadedDocument(
            content=content,
            filename=path.name,
            format=format_name,
            pages=pages,
            word_count=len(content.split()),
            char_count=len(content)
        )
    
    @staticmethod
    def _load_pdf(path: Path) -> Tuple[str, int]:
        """Load PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf required for PDF support. Run: pip install pypdf")
        
        reader = PdfReader(str(path))
        pages = len(reader.pages)
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {i+1} ---\n{text}")
        
        return "\n\n".join(text_parts), pages
    
    @staticmethod
    def _load_docx(path: Path) -> Tuple[str, int]:
        """Load Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for Word support. Run: pip install python-docx")
        
        doc = Document(str(path))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts), 1
    
    @staticmethod
    def _load_epub(path: Path) -> Tuple[str, int]:
        """Load EPUB file."""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("ebooklib and beautifulsoup4 required for EPUB. Run: pip install ebooklib beautifulsoup4")
        
        book = epub.read_epub(str(path))
        
        text_parts = []
        chapter_count = 0
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                chapter_count += 1
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                text = soup.get_text(separator='\n')
                if text.strip():
                    text_parts.append(f"--- Chapter {chapter_count} ---\n{text}")
        
        return "\n\n".join(text_parts), chapter_count
    
    @staticmethod
    def _load_html(path: Path) -> Tuple[str, int]:
        """Load HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 required for HTML. Run: pip install beautifulsoup4")
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        
        text = soup.get_text(separator='\n')
        
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        return "\n".join(lines), 1
    
    @staticmethod
    def _load_text(path: Path) -> Tuple[str, int]:
        """Load plain text file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return content, 1
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if a file format is supported."""
        ext = Path(file_path).suffix.lower()
        # We support all text-based files, not just the listed ones
        return True
    
    @classmethod
    def get_supported_formats(cls) -> str:
        """Get a string listing all supported formats."""
        formats = []
        for ext, name in cls.SUPPORTED_FORMATS.items():
            formats.append(f"{ext} ({name})")
        return ", ".join(formats)


def load_document(file_path: str) -> LoadedDocument:
    """
    Convenience function to load a document.
    
    Args:
        file_path: Path to the document
        
    Returns:
        LoadedDocument with extracted text
        
    Example:
        doc = load_document("manual.pdf")
        print(f"Loaded {doc.format}: {doc.word_count} words")
        print(doc.content)
    """
    return DocumentLoader.load(file_path)
